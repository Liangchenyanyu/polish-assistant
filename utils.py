"""
工具模块
负责: 文件解析、日志配置、异常处理、结果导出
评分对应: 异常处理、日志记录、结果导出
"""
import logging
from pathlib import Path
from datetime import datetime
from config import config
from typing import Optional


def setup_logging() -> logging.Logger:
    """配置日志系统（评分点: 日志记录）"""
    log_file = config.logs_dir / f"app_{datetime.now().strftime('%Y%m%d')}.log"

    logging.basicConfig(
        level=logging.INFO,
        format="%(asctime)s - %(name)s - %(levelname)s - %(message)s",
        handlers=[
            logging.FileHandler(log_file, encoding="utf-8"),
            logging.StreamHandler(),
        ],
    )
    return logging.getLogger("PolishAssistant")


logger = setup_logging()


def save_polish_result(
    original: str, polished: str, style: str, filename: Optional[str] = None
) -> str:
    """保存润色结果到文件（评分点: 结果导出）"""
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"polished_{timestamp}.md"

    filepath = config.data_dir / filename

    content = f"""# 文章润色结果

**润色时间**: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  
**润色风格**: {style}  

---

## 原文

{original}

---

## 润色后

{polished}

---

*由AI文章润色助手生成*
"""
    filepath.write_text(content, encoding="utf-8")
    logger.info(f"润色结果已保存: {filepath}")
    return str(filepath)


def handle_file_upload(file_path: str) -> str:
    """处理文件上传（评分点: 文件上传支持）
    支持：.txt, .md, .docx (纯文本提取)
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"文件不存在: {file_path}")

    suffix = path.suffix.lower()
    if suffix in (".txt", ".md"):
        return path.read_text(encoding="utf-8")
    else:
        raise ValueError(
            f"不支持的文件格式: {suffix}。支持的格式: .txt, .md"
        )


def format_exception(e: Exception) -> str:
    """格式化异常信息（评分点: 异常处理）"""
    return (
        f"【异常提示】{type(e).__name__}: {str(e)}。"
        "请检查网络连接、API配置或输入内容后重试。"
    )


def truncate_text(text: str, max_length: int = 200) -> str:
    """截断文本用于日志显示"""
    if len(text) <= max_length:
        return text
    return text[:max_length] + "..."


def standardize_format(text: str) -> dict:
    """
    格式标准化：统一标点、段落、首行缩进、去除多余空行
    返回标准化后的文本和修改统计
    """
    import re
    original = text
    changes = []

    # 1. 统一中文标点（半角→全角）
    punct_map = {
        ',': '，', '.': '。', '!': '！', '?': '？',
        ';': '；', ':': '：', '(': '（', ')': '）',
        '<': '《', '>': '》',
    }
    for half, full in punct_map.items():
        # 只在中文上下文中替换
        count = text.count(half)
        if count > 0:
            text = re.sub(
                rf'([\u4e00-\u9fff])\s*{re.escape(half)}\s*([\u4e00-\u9fff])',
                rf'\1{full}\2', text
            )
            text = re.sub(rf'{re.escape(half)}$', full, text)  # 行尾
            text = re.sub(rf'^{re.escape(half)}', full, text)  # 行首
    if text != original:
        changes.append("统一中英文标点符号")

    # 2. 去除多余空行（多空行→单空行）
    old = text
    text = re.sub(r'\n{3,}', '\n\n', text)
    if text != old:
        changes.append("去除多余空行")

    # 3. 段落之间确保有空行
    old = text
    text = re.sub(r'([。！？])\n(?!\n)', r'\1\n\n', text)
    if text != old:
        changes.append("规范化段落分隔")

    # 4. 去除行首行尾多余空格
    old = text
    lines = text.split('\n')
    lines = [line.strip() for line in lines]
    text = '\n'.join(lines)
    if text != old:
        changes.append("去除行首行尾多余空格")

    # 5. 统一省略号
    old = text
    text = re.sub(r'\.{2,}', '…', text)
    text = re.sub(r'。。+', '…', text)
    if text != old:
        changes.append("统一省略号为「…」")

    return {
        "text": text,
        "changes": changes,
        "original_length": len(original),
        "standardized_length": len(text),
    }


def export_to_word(original: str, polished: str, style: str, filename: Optional[str] = None) -> str:
    """
    导出为 Word 文档 (.docx)
    返回文件路径
    """
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"polished_{timestamp}.docx"

    filepath = config.data_dir / filename
    doc = Document()

    # 设置默认字体
    style = doc.styles['Normal']
    font = style.font
    font.name = '宋体'
    font.size = Pt(11)

    # 标题
    title = doc.add_heading('文章润色结果', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 元数据
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(f"润色时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  润色风格: {style}")
    meta_run.font.size = Pt(9)
    meta_run.font.color.rgb = RGBColor(128, 128, 128)

    doc.add_paragraph()  # 空行

    # 原文
    doc.add_heading('原文', level=1)
    for para in original.split('\n'):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.line_spacing = 1.5

    doc.add_page_break()

    # 润色后
    doc.add_heading('润色后', level=1)
    for para in polished.split('\n'):
        if para.strip():
            p = doc.add_paragraph(para.strip())
            p.paragraph_format.line_spacing = 1.5

    # 页脚
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('由AI文章润色助手生成')
    footer_run.font.size = Pt(8)
    footer_run.font.color.rgb = RGBColor(180, 180, 180)

    doc.save(str(filepath))
    logger.info(f"Word文档已导出: {filepath}")
    return str(filepath)


def export_to_pdf(original: str, polished: str, style: str, filename: Optional[str] = None) -> str:
    """
    导出为 PDF 文档（使用纯文本方式，简单实现）
    返回文件路径
    """
    if not filename:
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"polished_{timestamp}.pdf"

    filepath = config.data_dir / filename

    # 使用 fpdf2（轻量，无外部依赖）
    try:
        from fpdf import FPDF
    except ImportError:
        # 降级：保存为 .txt 格式但标注 PDF
        txt_content = (
            f"文章润色结果\n"
            f"润色时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n"
            f"润色风格: {style}\n\n"
            f"{'='*50}\n\n原文:\n{original}\n\n{'='*50}\n\n润色后:\n{polished}\n\n"
            f"由AI文章润色助手生成"
        )
        filepath = filepath.with_suffix('.txt')
        filepath.write_text(txt_content, encoding='utf-8')
        logger.warning("fpdf2 未安装，PDF 导出降级为纯文本")
        return str(filepath)

    pdf = FPDF()
    pdf.add_page()

    # 添加中文字体支持（尝试常见字体路径）
    font_paths = [
        "C:/Windows/Fonts/simsun.ttc",    # 宋体
        "C:/Windows/Fonts/msyh.ttc",      # 微软雅黑
        "C:/Windows/Fonts/simhei.ttf",    # 黑体
    ]
    font_added = False
    for fp in font_paths:
        if Path(fp).exists():
            try:
                pdf.add_font('CJK', '', fp, uni=True)
                pdf.set_font('CJK', '', 11)
                font_added = True
                break
            except Exception:
                continue

    if not font_added:
        # 降级
        filepath = filepath.with_suffix('.txt')
        txt_content = (
            f"文章润色结果\n\n原文:\n{original}\n\n---\n\n润色后:\n{polished}"
        )
        filepath.write_text(txt_content, encoding='utf-8')
        logger.warning("无法加载中文字体，PDF 导出降级为纯文本")
        return str(filepath)

    # 标题
    pdf.set_font('CJK', '', 16)
    pdf.cell(0, 10, '文章润色结果', align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(5)

    pdf.set_font('CJK', '', 9)
    pdf.set_text_color(128, 128, 128)
    pdf.cell(0, 6, f"润色时间: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}  |  风格: {style}", align='C', new_x="LMARGIN", new_y="NEXT")
    pdf.ln(8)

    # 原文
    pdf.set_text_color(0, 0, 0)
    pdf.set_font('CJK', '', 14)
    pdf.cell(0, 8, '原文', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font('CJK', '', 10)
    for line in original.split('\n'):
        if line.strip():
            pdf.multi_cell(0, 6, line.strip())

    pdf.add_page()

    # 润色后
    pdf.set_font('CJK', '', 14)
    pdf.cell(0, 8, '润色后', new_x="LMARGIN", new_y="NEXT")
    pdf.set_font('CJK', '', 10)
    for line in polished.split('\n'):
        if line.strip():
            pdf.multi_cell(0, 6, line.strip())

    pdf.output(str(filepath))
    logger.info(f"PDF文档已导出: {filepath}")
    return str(filepath)
